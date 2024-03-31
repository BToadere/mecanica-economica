# -*- coding: utf-8 -*-
"""
Created on Wed Dec 20 20:48:29 2023

@author: bogdan.stefan
"""

# from IPython import get_ipython
# get_ipython().magic('clear')

# import sys

import pandas as pd
import numpy as np
import copy
from scipy.interpolate import *
import matplotlib.pyplot as plt

import os
import openpyxl
import docx
from docx.shared import Pt, Mm
import subprocess
import datetime


# =============================================================================
# ~ Funciones para exportar los datos a excel.
# =============================================================================

                   
def export_data_frame_to_excel(ruta_excel, lista_tablas, num_imagenes):
    
    fila = int(num_imagenes/2)*7+2
    print(ruta_excel)
    # Crear un escritor de Excel fuera del bucle
    with pd.ExcelWriter(ruta_excel, engine='openpyxl') as writer:
        workbook = writer.book
        
        for tabla in lista_tablas:
            nombre_tabla = [nombre for nombre, valor in globals().items() if valor is tabla][0]
            # print("Exportando la tabla:", nombre_tabla)
            # try:
            # Escribir la tabla en el escritor de Excel
            tabla.to_excel(writer, sheet_name=nombre_tabla, startrow=fila, startcol=0, index=False)
            sheet = workbook[nombre_tabla]
            
            # Configurar el formato de la celda (fuente Calibri 9)
            # fuente = Font(name= 'Calibri', size=9)

            for column in sheet.columns:
                # print('columna', column)
                max_length = 0
                column = [cell for cell in column]
                # print('columna', column)
                for cell in column:
                    # print('\ncelda:', cell.column_letter)
                    # cell.font = openpyxl.styles.Font(name='Calibri', size=8.5)
                    try:
                        if len(str(cell.value)) > max_length:
                            max_length = len(str(cell.value))
                            # print(max_length)
                    except:
                        pass
                adjusted_width = (max_length + 2)
                # print(adjusted_width)
                # print(adjusted_width)
                # sheet.pageSetUpPr = PageSetupProperties(fitToPage=True, autoPageBreaks=True)
                # workbook.save()
                # print(cell.column_letter, sheet.column_dimensions[cell.column_letter].width)
                
                
                # workbook.save('DATAFRAME.xlsx')
                # workbook.close()
                # wb = xw.Book(ruta_excel)
                # ws = wb.sheets[nombre_tabla]
                # ws.autofit()
                # wb.save('DATAFRAME.xlsx')
                # wb.close()
            # except Exception as e:
            #     print(f"Error al exportar la tabla {nombre_tabla}: {e}")
            
            #CALIBRI 9
        
                   
def export_plots_to_excel(carpeta_imagenes, nombre_archivo_excel):
    # Abre archivo de Excel
    wb = openpyxl.load_workbook(nombre_archivo_excel)
               
    # Crear una hoja llamada "Gráficas"
    # hoja = wb.create_sheet(title="Gráficas")
    hoja = wb['se_z']

    # Obtener la lista de nombres de archivos en la carpeta de imágenes
    archivos = sorted(os.listdir(carpeta_imagenes), key=lambda x: os.path.getctime(os.path.join(carpeta_imagenes, x)))
    # print(archivos)
    nombres_imagenes = [archivo for archivo in archivos if archivo.lower().endswith('.png')]
    # print(nombres_imagenes)
    
    # Escribir las imágenes en la hoja de Excel
    for indice, nombre_imagen in enumerate(nombres_imagenes, start=1):
        # print(nombre_imagen)
        ruta_imagen = os.path.join(carpeta_imagenes, nombre_imagen)
        # print(ruta_imagen)
        imagen = Image(ruta_imagen)

        # Establecer el tamaño de la imagen (hoja A4 10*44 celdas)
        celda_largo = 64
        celda_alto = 19
        imagen.width = celda_largo * 5
        imagen.height = celda_alto * 7
        
        # Calcular la celda en la que se insertará la imagen
        # Calcular la celda en la que se insertará la imagen
        if indice % 2 == 1:
            columna = 'A'
            if indice == 1:
                fila = 1
            else:
                fila += 7  # Ajusta según sea necesario
        else:
            columna = 'F'
        celda = f'{columna}{fila}'
        # print(indice, celda)
        # celda = f'A{(indice-1) * 15+1}'  # Aumentar el índice multiplicador según sea necesario

        # Insertar la imagen en la hoja de Excel
        hoja.add_image(imagen, celda)
        
    #Maketación, es para qeu la oclumna E sea mas estrecha y se imprima bien
    for column in hoja.columns:
        max_length = 0
        column_letter = column[0].column_letter
        for cell in column:
            try:
                if len(str(cell.value)) > max_length:
                    max_length = len(cell.value)
            except:
                pass
        adjusted_width = (max_length + 2) * 1.2
        hoja.column_dimensions[column_letter].width = adjusted_width
        
    hoja.column_dimensions['A'].width = 6.2
    # Guardar el archivo de Excel
    wb.save(nombre_archivo_excel)
    
    # # Descomentar para borrar carpeta de imagenes
    # # Proceso para reiniciar la tabla auxiliar de imagenes
    # try:
    #     shutil.rmtree(carpeta_imagenes)
    #     print(f"Carpeta '{carpeta_imagenes}' eliminada con éxito.")
    # except OSError as e:
    #     print(f"No se pudo eliminar la carpeta '{carpeta_imagenes}': {e}")


def exportar_series(lista_tablas, ruta_excel, carpeta_con_imagenes):
    try:
        # Contar la cantidad de imagenes
        cantidad_imagenes = len(os.listdir(carpeta_con_imagenes))
        export_data_frame_to_excel(ruta_excel, lista_tablas, cantidad_imagenes)
        # print(carpeta_con_imagenes)
        # export_plots_to_excel(carpeta_con_imagenes, ruta_excel)
    except PermissionError:
        print('Error: El archivo está ABIERTO. Por favor, cierra el archivo y vuelve a intentarlo.')
    
    # Proceso para abrir el excel despues de exportar
    try:
        subprocess.Popen(['start', 'excel', ruta_excel], shell=True)
    except FileNotFoundError:
        print(f"Error: El archivo '{ruta_excel}' no fue encontrado.")
    except openpyxl.utils.exceptions.InvalidFileException:
        print(f"Error: '{ruta_excel}' no es un archivo Excel válido.")
    except Exception as e:
        print(f"Error al abrir el archivo '{ruta_excel}': {e}")


# =============================================================================
# ~ Función para graficar, el atributo es UNA LISTA DE TABLAS, si quieres pinta 
# otra columna cambia el nombre donde he señalado
# =============================================================================
def pinta(tablas, columnas):
    # print('_____________________________________________________________\n_____________________________________________________________\n\n')
    # Crear la carpeta 'Exportaciones' si no existe
    carpeta_exportaciones = 'temp_graficas'
    if not os.path.exists(carpeta_exportaciones):
        os.makedirs(carpeta_exportaciones)
    
    # Inicializar la figura
    plt.figure()    
    
    for i in range(len(tablas)):
        tabla = tablas[i]
        # print(tabla)
        # print('tabla:\n', tabla)
        # Extraer el nombre de la tabla
        nombre_tabla = [nombre for nombre, valor in globals().items() if valor is tabla][0]
        # print(nombre_tabla, columnas)
        # Crear la ruta completa para guardar la figura
        ruta_figura = os.path.join(carpeta_exportaciones, f'{nombre_tabla}_{"_".join(columnas[i])}.png')
        # print(ruta_figura)
        
        
        # Iterar sobre las columnas
        # columna_dominio = columnas[0][0]
        # print('inicio: ', columna_dominio, len(tabla[columna_dominio]), '\n', tabla[columna_dominio])
        # for columna in columnas[i]:
        #     print('candidato: ', columna, len(columna))
        #     print('esto queremso ocmprobar', len(tabla[columna]), ' > ', len(tabla[columna_dominio]))
        #     if len(tabla[columna]) > len(tabla[columna_dominio]):
        #         print('entra porque '. len(tabla[columna]), ' > ', len(tabla[columna_dominio]))
        #         columna_dominio = columna
        # print(columna_dominio)
        
        for columna in columnas[i]:
            # print('columna:\n', columna)
            # Extraer datos de la tabla
            x = tabla['TIME'].apply(lambda x: x.timestamp()).to_numpy()
            y = tabla[columna].to_numpy()  # Columna a graficar
            # print(type(x), type(y))
            # print(len(x), len(y))
    
            # Eliminar elementos nulos
            x = x[~np.isnan(y)]
            y = y[~np.isnan(y)]
            # print(len(x), len(y))
            # print(x, y)
            if len(y) == 0:
                print(f'No se ha podido graficar la {nombre_tabla}_{columna}')
                continue
    
            dom_ini = x[0]
            dom_fin = x[-1]
            
            # Generar nuevos puntos para la interpolación
            xs = np.linspace(dom_ini, dom_fin, 100)
            fec_x = [datetime.datetime.fromtimestamp(ts) for ts in x]
            fec_xs = [datetime.datetime.fromtimestamp(ts) for ts in xs]

            # print(fec_xs)
            
            try:    
                # Intentar la interpolación cuadrática
                f_interp = interp1d(x, y, kind='quadratic')     # Puedes cambiar 'linear' a 'cubic', 'quadratic', etc.
                ys_interp = f_interp(xs)
                ys_interp = f_interp(xs)
            except ValueError:
                # En caso de error, realizar la interpolación lineal
                # print(f'Interpolación cuadrática fallida para {nombre_tabla}_{columna}. Realizando interpolación lineal.')
                f_interp = interp1d(x, y, kind='linear')
                ys_interp = f_interp(xs)
            # # Interpolación polinómica usando el método scipy
            # f_interp = interp1d(x, y, kind='cuadratic')  
    
            # Graficar los resultados
            plt.plot(fec_x, y, 'o')#, label=f'Data- {columna}')         #creo que hay qeu cambiarlo aqui
            plt.plot(fec_xs, ys_interp, label=f'{columna}')
        
    # Personalizar ejes
    # plt.xlabel('Eje X')  # Nombre del eje X
    # plt.ylabel('Eje Y')  # Nombre del eje Y
    plt.title('\n'.join(columnas[i]))                  # Título de la gráfica
    # plt.axhline(0, color='black',linewidth=0.5)     # Línea horizontal en y=0
    # plt.axvline(25, color='black',linewidth=0.5)     # Línea vertical en x=0
    plt.grid(axis='y', linestyle='--', alpha=0.7, which='both')  # Agregar malla horizontal
    
    # Modificar límites de los ejes
    # plt.xlim(-2*np.pi, 2*np.pi)
    # plt.ylim(-1.5, 1.5)
    
    # Añadir leyenda
    plt.legend()
    
    # Guardar figura
    plt.savefig(ruta_figura)
    
    # Mostrar la gráfica
    # plt.show()
    


# =============================================================================
# ~ Función para calcular la volatilidad
# =============================================================================
def volatilidad_vec(columna, grupo):
    N = len(columna)
    result = np.empty(N)
    result[:] = np.nan
    if N < grupo:
        print('No se ha podido calcular la volatilidad en el rango de ', 
                    grupo, ' hay solamente ', 
                    N, ' elementos.')
        return result

    if not isinstance(columna, np.ndarray):
        columna = columna.to_numpy()

    # Aplicar logaritmo a la columna
    columna = np.log(columna)

    # Calcular la diferencia entre elementos consecutivos (analogo formula articulo)
    columna = np.diff(columna)#np.insert(, 0, 0)
    # print(N, len(columna))
    
    # print(result)
    
    for i in range(grupo, N):
        # print(i, range(i-grupo, i))
        result[i] = np.std(columna[i-grupo:i+grupo])
    # print(len(result), result)
    return result

def entropia_shannon(valor):
    max_valor = np.max(valor)
    normal = valor/max_valor
    # print(es_max['valor'].iloc[0])
    orden = 10**int(np.log10(max_valor))
    # print(es_max)
    # print(orden)
    normal = normal*orden
    normal = np.round(normal)
    
    # print(normal)
    N = len(normal)
    
    result = np.empty(N)
    result[:] = np.nan
    for i in range(N):
        cantidad = np.count_nonzero(normal == normal[i])
        result[i] = cantidad
    
    frecuencia = result / N
    entropia = np.log(frecuencia)*frecuencia
    # print(entropia)
    
    # tabla_x['ENTROPY'] = entropia
    
    # Resto de la lógica de tu función aquí
    
    return entropia
        
    

# =============================================================================
# ~ Función para calcular serie X
# =============================================================================
def serie_X(selfa, val_p, grupo, val_r):
    rango_datos = len(selfa) - val_r
    valor = selfa['valor'].to_numpy()

    # Normalizar la serie de valores
    normal = valor / np.max(valor[:rango_datos])

    # Calcular la variación absoluta y la normalización filtrada
    var_abs = np.insert(np.abs(np.diff(normal)), 0, 0)
    norm_filtr = copy.copy(normal)

    # Filtrar valores según un umbral (val_p)
    var_abs_filtr = copy.copy(var_abs)
    ind_valp = var_abs >= val_p
    norm_filtr[~ind_valp] = 0
    var_abs_filtr[~ind_valp] = 0

    # Calcular volatilidad
    volatilidad = volatilidad_vec(normal, grupo)
    entropia = np.empty(len(valor))
    entropia[:] = np.nan
    entropia[:rango_datos] = entropia_shannon(valor[:rango_datos])
    # Agregar nuevas columnas al DataFrame
    selfa['INITIAL SERIE'] = normal
    selfa['var_abs'] = var_abs
    selfa['ind_valp'] = ind_valp
    selfa['norm_valp'] = norm_filtr
    selfa['var_valp'] = var_abs_filtr
    selfa['VOLATILITY'] = volatilidad
    selfa['ENTROPY'] = entropia

# =============================================================================
# ~ Función para calcular serie Y
# =============================================================================
def serie_Y(tabla, val_r):
    # tabla.head(len(tabla) - val_r)
    grupo = 20  # Valor a agrupar volatilidad
    serie_y = pd.DataFrame()

    # Extraer columnas de la tabla
    tiempo = tabla['TIME'].to_numpy()
    normal = tabla['INITIAL SERIE'].to_numpy()
    var_abs = tabla['var_abs'].to_numpy()
    ind_valp = tabla['ind_valp'].to_numpy()

    # Filtrar valores según el indicador ind_valp
    temp_filtr = tiempo[ind_valp]
    norm_filtr = normal[ind_valp]
    var_filtr = var_abs[ind_valp]

    # Calcular volatilidad
    volatilidad = volatilidad_vec(norm_filtr, grupo)

    # Crear DataFrame de la serie Y
    serie_y['TIME'] = temp_filtr
    serie_y['normal'] = norm_filtr
    serie_y['var_abs'] = var_filtr
    # serie_y['volatilidad'] = volatilidad
    return serie_y

# =============================================================================
# ~ Función para calcular serie Z
# =============================================================================
def serie_Z(tabla_x, val_r, boton):
    grupo = 3 # Valor a agrupar volatilidad
    serie_z = pd.DataFrame()
    
    promedio_prediccion = np.mean(tabla_x[len(tabla_x) - val_r:]['INITIAL SERIE'])
    # print(tabla_x[len(tabla_x) - val_r:])
    # print(promedio_prediccion)
    nuevo_rango = len(tabla_x)
    nuevo_tiempo = tabla_x['TIME'][nuevo_rango-1]
    # Truncamos la tabla para hacer los calculos sobre lso datos deseados
    tabla_x = tabla_x.head(len(tabla_x) - val_r)
    # Extraer columnas de la tabla_x
    tiempo_original = list(tabla_x['TIME'])
    val_rang = tabla_x['norm_valp'].to_numpy()
    var_rang = tabla_x['var_valp'].to_numpy()
    ind_ini = tabla_x['ind_valp'].to_numpy()
    N = len(ind_ini)

    # Ajustar el valor de val_r si el botón es 1
    if boton == 1:
        val_r = int(N / val_r)

    sobran = N % val_r
    if sobran >= 0:
        val_rang = val_rang[sobran:]
        var_rang = var_rang[sobran:]
        ind_ini = ind_ini[sobran:]

    # Reshape para dividir en grupos de val_r
    val_rang = val_rang.reshape(-1, val_r)
    ind_ini = ind_ini.reshape(-1, val_r)

    # Contar la cantidad de unos en cada grupo
    count_nodo = np.sum(ind_ini, axis=1)

    # Calcular densidad, diferencias y promedios
    densidad = count_nodo / np.max(count_nodo)
    den_diff = np.insert(np.diff(densidad), 0, None)
    val_prom = np.mean(val_rang, axis=1)
    val_diff = np.diff(val_prom)
    val_diff2 = np.insert(np.diff(val_diff), 0, None)

    # Añadir NaN's por el tamaño
    val_diff = np.insert(val_diff, 0, None)
    val_diff2 = np.insert(val_diff2, 0, None)

    # Calcular velocidad y sus diferencias
    velocidad = densidad * val_diff
    vel_diff = np.diff(velocidad)
    vel_diff2 = np.insert(np.diff(vel_diff), 0, None)

    # Añadir NaN's por el tamaño
    vel_diff = np.insert(vel_diff, 0, None)
    vel_diff2 = np.insert(vel_diff2, 0, None)

    # Calcular presión y sus diferencias
    presion = -densidad * velocidad ** 2
    pres_diff = np.insert(np.diff(presion), 0, None)

    # Calcular término de viscosidad y viscosidad
    viscosidad = np.empty(len(densidad))
    term_visc = np.empty(len(densidad))
    viscosidad_ABE = np.empty(len(densidad))
    viscosidad[:] = np.nan
    term_visc[:] = np.nan
    viscosidad_ABE[:] = np.nan
    m = 1
    for i in range(0, len(densidad)):
        if presion[i]*val_diff[i] != 0:
            term_visc[i] = velocidad[i]*(vel_diff[i]/val_diff[i])+(m/densidad[i])*((pres_diff[i]/val_diff[i])-(presion[i]/densidad[i])*(den_diff[i]/val_diff[i]))
            # print(visc_comun)
            if i + 1 < len(densidad): # and visc_comun != 0
                viscosidad[i] = ((velocidad[i+1]-velocidad[i])+velocidad[i]*(vel_diff[i]/val_diff[i])+(m/densidad[i])*((pres_diff[i]/val_diff[i])-(presion[i]/densidad[i])*(den_diff[i]/val_diff[i])))*densidad[i]/(((2*(val_diff[i-1]*velocidad[i]-(val_diff[i-1]+val_diff[i-2])*velocidad[i-1]+val_diff[i-2]*velocidad[i-2])/(val_diff[i-1]*val_diff[i-2]*(val_diff[i-1]+val_diff[i-2])))))
            
            viscosidad_ABE[i] = term_visc[i]*densidad[i]/((2*(val_diff[i-1]*velocidad[i]-(val_diff[i-1]+val_diff[i-2])*velocidad[i-1]+val_diff[i-2]*velocidad[i-2])/(val_diff[i-1]*val_diff[i-2]*(val_diff[i-1]+val_diff[i-2]))))

    # Diferencia entre viscosidad y viscosidad ABE
    substrac_visco = viscosidad_ABE - viscosidad
    # Calcular volatilidad
    volatilidad = volatilidad_vec(val_prom, grupo)
    
    # Crear DataFrame de la serie Z
    rangos = list(range(val_r, N + 1, val_r))
    # print(rangos)
    serie_z['TIME'] = [tiempo_original[x-1] for x in rangos]
    serie_z['RANGE'] = rangos #[f"{i}-{i+25}" for i in range(0, N, val_r)]
    # print(len(serie_z), len(densidad))
    serie_z['DENSITY'] = densidad
    serie_z['den_diff'] = den_diff
    serie_z['MEAN_VALUE'] = val_prom
    serie_z['val_diff'] = val_diff
    # serie_z['val_diff2'] = val_diff2
    serie_z['VELOCITY'] = velocidad
    serie_z['vel_diff'] = vel_diff
    # serie_z['vel_diff2'] = vel_diff2
    serie_z['PRESSURE'] = presion
    serie_z['pres_diff'] = pres_diff
    serie_z['TERM_VISC'] = term_visc
    serie_z['VISCOSITY'] = viscosidad
    serie_z['VISCOSITY ABE'] = viscosidad_ABE
    serie_z['SUBTRACTION VISCOSITIES'] = substrac_visco
    # serie_z['volatilidad'] = volatilidad
    
    nueva_fila = serie_z.iloc[0].copy()
    
    nueva_fila['TIME'] = nuevo_tiempo
    nueva_fila['RANGE'] = nuevo_rango
    nueva_fila['DENSITY'] = np.nan
    nueva_fila['MEAN_VALUE'] = promedio_prediccion
    # Agregar la nueva fila al DataFrame
    serie_z.loc[len(serie_z['RANGE'])] = nueva_fila
    
    return serie_z

def prediccion_velocidad(serie_z, seed_point = 0, seed_distance = 0):
    
    serie_z_predic = serie_z.copy()
    # print('COPIA: ', serie_z_predic)
    # print(len(serie_z_predic), len(densidad))
    densidad = serie_z_predic['DENSITY'].to_numpy()
    den_diff = serie_z_predic['den_diff'].to_numpy() 
    # val_prom = serie_z_predic['MEAN_VALUE'].to_numpy()
    val_diff = serie_z_predic['val_diff'].to_numpy()
    # val_diff2 = serie_z_predic['val_diff2'].to_numpy()
    velocidad = serie_z_predic['VELOCITY'].to_numpy()
    vel_diff = serie_z_predic['vel_diff'].to_numpy() 
    # vel_diff2 = serie_z_predic['vel_diff2' ].to_numpy()
    presion = serie_z_predic['PRESSURE'].to_numpy()
    pres_diff = serie_z_predic['pres_diff'].to_numpy()
    # term_visc = serie_z_predic['TERM_VISC'].to_numpy()
    viscosidad = serie_z_predic['VISCOSITY'].to_numpy()
    
    # print('CONTROL: ', serie_z)
    t = len(viscosidad)-2
    # print(t)
    # print( nueva_fila)
    viscosidad[t] = seed_point + seed_distance
    # print('vsico ini: ', viscosidad[t])
    
    # print('ojo: ', serie_z_predic) #serie_z_predic modificada
    # print(viscosidad)
    # print(t)
    # serie_z_predic.iloc[t+1]['VELOCITY'] 
    prediccion= velocidad[t]-velocidad[t]*(vel_diff[t]/val_diff[t])-(1/densidad[t])*((pres_diff[t]/val_diff[t])-(presion[t]/densidad[t])*(den_diff[t]/val_diff[t]))+(viscosidad[t]/densidad[t])*((2*(val_diff[t-1]*velocidad[t]-(val_diff[t-1]+val_diff[t-2])*velocidad[t-1]+val_diff[t-2]*velocidad[t-2])/(val_diff[t-1]*val_diff[t-2]*(val_diff[t-1]+val_diff[t-2]))))
    serie_z_predic['VELOCITY'].iloc[-1] = prediccion
    serie_z_predic['val_diff'].iloc[-1] = serie_z_predic['MEAN_VALUE'].iloc[-1] - serie_z_predic['MEAN_VALUE'].iloc[-2]
    serie_z_predic['DENSITY'].iloc[-1] = serie_z_predic['VELOCITY'].iloc[-1]/serie_z_predic['val_diff'].iloc[-1]
    serie_z_predic['SUBTRACTION VISCOSITIES'].iloc[-2] = -seed_distance
    
    serie_z_predic = serie_z_predic.tail(2)
    serie_z_predic.rename(columns={'VELOCITY': 'VELOCITY PREDICT'}, inplace=True)
    # print('prediccion: \n', serie_z_predic.to_string())
    
    
    
    
    
    
    
    return serie_z_predic

def calculos_resumen_serie(serie_x, serie_z, val_p, val_r, seeds):
    tabla_resumen = pd.DataFrame()
    # print(serie_x)
    velocidad = serie_z['VELOCITY'].to_numpy()
    viscosidad = serie_z['VISCOSITY'].to_numpy()
    viscosidad_ABE = serie_z['VISCOSITY ABE'].to_numpy()
    substrac_visco  = serie_z['SUBTRACTION VISCOSITIES'].to_numpy()
    
    entropy = serie_x['ENTROPY'].to_numpy()
    ind_valp = serie_x['ind_valp'].to_numpy()
    # CALCULO IC
    max_aux = np.maximum(viscosidad_ABE, viscosidad)
    min_aux = np.minimum(viscosidad_ABE, viscosidad)
    IC = (np.nansum(abs(substrac_visco))/len(substrac_visco))/(
        np.nanmax(max_aux)- np.nanmin(min_aux))
    # print('IC: ', IC)
    
    # CALCULO SHANNON ENTROPY
    shannon_entropy = abs(np.sum(entropy))
    # print('SHANNON ENTROPY: ', shannon_entropy)
    
    # CALCULO SUBSTRACTION VISCOSITIES
    max_substract_visco = np.nanmax(substrac_visco)
    min_substract_visco = np.nanmin(substrac_visco)
    diff_max_min_visco = max_substract_visco - min_substract_visco
    # print(diff_max_min_visco, max_substract_visco, min_substract_visco)
    
    #CALULO TYPICAL DEVIATION VELOCITY
    std_velocity = np.nanstd(velocidad, ddof=1)
    # print('DESVIACION TIP VELOCIDAD: ', std_velocity)
    
    # CALCULO NODOS
    num_points = len(ind_valp)
    num_nodes = np.count_nonzero(ind_valp)
    # print(num_points, val_p, num_nodes)
    
    # CALCULOS AVERAGE SUBSTRAC VISCOSITIES
    mean_substrc_visco = np.nanmean(substrac_visco)
    
    # Calculo NUMERO DE INERVALOS
    num_intervalos = len(velocidad)
    
    tabla_resumen['IC'] = [IC]
    tabla_resumen['SHANNON ENTROPY'] = [shannon_entropy] 
    tabla_resumen['MAX-MIN (SUBSTRACTION VISCOSITIES)'] = [diff_max_min_visco]
    tabla_resumen['MAX (SUBSTRACTION VISCOSITIES)'] = [max_substract_visco]
    tabla_resumen['MIN (SUBSTRACTION VISCOSITIES)'] = [min_substract_visco]
    tabla_resumen['IS (STANDARD DEVIATION VELOCITY)'] = [std_velocity]
    tabla_resumen['AVERAGE SUBSTRACION VISCOSITIES'] = [mean_substrc_visco]
    tabla_resumen['SERIE NUMBER TOTAL'] = [num_points]
    tabla_resumen['\"p\" VALUE FOR NODES'] = [val_p]
    tabla_resumen['NODES NUMBER'] = [num_nodes]
    tabla_resumen['INTERVAL NUMBER'] = [num_intervalos]
    tabla_resumen['NUMBER POINTS EACH INTERVAL'] = [val_r]
    
    for i, seed in enumerate(seeds, 1):
        nom_columna = f'SEED - {i}'
        tabla_resumen[nom_columna] = seed
    
    print('RESUMEN\n\n', tabla_resumen.to_string())
    return tabla_resumen

def ajustar_margenes(doc, margen_izquierdo=1, margen_derecho=1, margen_superior=1, margen_inferior=1):
    sections = doc.sections
    for section in sections:
        section.left_margin = Mm(margen_izquierdo)
        section.right_margin = Mm(margen_derecho)
        section.top_margin = Mm(margen_superior)
        section.bottom_margin = Mm(margen_inferior)
    
def cambiar_interlineado(parrafo, interlineado):
    for run in parrafo.runs:
        run.font.size = Pt(7)  # Ajusta el tamaño de fuente según sea necesario
        run._element.get_or_add_pPr().get_or_add_spacing().line = Pt(interlineado)
    
def ajustar_interlineado_entre_parrgrafos(parrafo, espacio_despues):
    parrafo.paragraph_format.space_after = Pt(espacio_despues)


# def exportar_tabla_word(doc, tabla):
#     # Initialise the table
#     t = doc.add_table(rows=1, cols=tabla.shape[1])
#     # Add the column headings
#     for j in range(tabla.shape[1]):
#         cell = tabla.columns[j].strip()
#         p = t.cell(0, j).add_paragraph()
#         # ajustar_interlineado_entre_parrgrafos(p, 0.2)
#         p.add_run(str(cell)).bold = True
#         # print(str(cell))
#     # Add the body of the data frame
#     for i in range(tabla.shape[0]):
#         row = t.add_row()
#         for j in range(tabla.shape[1]):
#             cell = tabla.iat[i, j]
#             row.cells[j].text = str(cell)

def exportar_tabla_word(doc, tabla):
    # Initialise the table
    t = doc.add_table(rows=tabla.shape[1], cols=2)
    # t.style = 'List Table 2 Accent 1'
    t.autofit = True
    # print(type(t))
    # Add the column headings
    for j in range(tabla.shape[1]):
        titulo = tabla.columns[j].strip()
        valor = tabla.iloc[(0,j)]
        t.cell(j, 0).text = str(titulo)
        t.cell(j, 0).paragraphs[0].runs[0].bold = True
        t.cell(j, 1).text = str(valor)
    t.autofit = True
            
def exportar_tabla_word_bullet (doc, tabla):
        for columna in tabla.columns:
            p = doc.add_paragraph(style = 'List Bullet')
            # Añade el texto del título de la columna en negrita
            p.add_run(columna.strip() + ': ').bold = True
            # Añade el valor de la posición 0 sin negrita
            p.add_run(str(tabla[columna][0])).bold = False
            
        doc.add_paragraph()
        
def exportar_tabla_word_t(doc, tabla, transpuesto=False):
    # Initialise the table
    if transpuesto:
        t = doc.add_table(rows=tabla.shape[1]+1, cols=tabla.shape[0])
    else:
        t = doc.add_table(rows=1, cols=tabla.shape[1])

    # Add the column headings
    if transpuesto:
        for i in range(tabla.shape[1]):
            cell = tabla.index[i]
            t.cell(i+1, 0).text = str(cell)
            t.cell(i+1, 0).paragraphs[0].runs[0].bold = True
    else:
        for j in range(tabla.shape[1]):
            cell = tabla.columns[j].strip()
            t.cell(0, j).text = str(cell)
            t.cell(0, j).paragraphs[0].runs[0].bold = True

    # Add the body of the data frame
    for i in range(tabla.shape[0]):
        for j in range(tabla.shape[1]):
            if transpuesto:
                cell = tabla.iat[i, j]
                t.cell(j+1, i).text = str(cell)
            else:
                cell = tabla.iat[i, j]
                t.cell(i+1, j).text = str(cell)

            
def exportar_plots_word(doc, carpeta_con_imagenes):
    lista_imagenes = [archivo for archivo in os.listdir(carpeta_con_imagenes) if archivo.lower().endswith('.png')]
    paragraph = doc.add_paragraph()
    ajustar_interlineado_entre_parrgrafos(paragraph, 0.2)
    for i, imagen in enumerate(lista_imagenes, 1):
        ruta_imagen = os.path.join(carpeta_con_imagenes, imagen)
        if i % 3+ 1:
            paragraph.add_run().add_text('  ')
        run_imagen = paragraph.add_run()
        run_imagen.add_picture(ruta_imagen, width=Mm(63))  # Ajusta el ancho según sea necesario
        if i%3 == 0:
            paragraph = doc.add_paragraph()
            ajustar_interlineado_entre_parrgrafos(paragraph, 0.2)
    return paragraph
            
def export_ficha_word(resumen_caracteristicas, serie_z, carpeta_con_imagenes, nombre_datos=''):
    
    doc = docx.Document()
    
    ajustar_margenes(doc, 10, 5, 0, 0)
    
    doc.add_heading(f'Ficha: {nombre_datos}', 0)
    
    # p = doc.add_paragraph('The ')
    # p.add_run('International System of Units').bold = True
    # p.add_run(', known by the international abbreviation ')
    # p.add_run('SI').bold = True
    # p.add_run(', is the modern form of the metric system ')
    # p.add_run('(Wikipedia: International System of Units)').italic = True
    
    # doc.add_heading('Properites', level=1)
    exportar_tabla_word(doc, resumen_caracteristicas)
    doc.add_paragraph('')
    
    parrafo = exportar_plots_word(doc, carpeta_con_imagenes)
    ajustar_interlineado_entre_parrgrafos(parrafo, 5)       
    # doc.add_heading('Table Serie Z', level=1)
    # exportar_tabla_word(doc, serie_z)
    
    # Settings
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Verdana'
    font.size = Pt(9)
    
    # Save the Word doc
    doc.save('Ficha_Serie.docx') 
    
    
    
    
    
    
    
# =============================================================================
# ~                                 CODIGO
# =============================================================================

# # Ruta para PERSONAL
# ruta = r"G:\My Drive\A.Puente_SDG\CFD - Timoteo\Python\Data\Datos_Fase1.xlsx"

# Ruta para SDG
ruta = './data/Datos_Fase1.xlsx'
# Local
# ruta = r"C:\Users\bogdan.stefan\Desktop\Temp_LOCAL\Datos_T1.xlsx"


hoja= 'datos_it4 (fecha)'

cantidad_datos = 200

val_p = 0.002                   # Valor para discriminar NODOS
grupo_volatilidad = 20          # Valor a agrupar volatilidad

val_r = 25                      # Valor para agrupar los datos en paquetes
boton = 0                       # 0 indica que val_r es el num de elementos de cada subgrupo, 
                                # 1 indica que val_r es el num de subgrupos. 
datos_para_prediccion = cantidad_datos + val_r

se_x = pd.read_excel(ruta, sheet_name=hoja)
se_x.columns= ['TIME', 'valor']

if cantidad_datos > 0:
    se_x = se_x.head(datos_para_prediccion)


serie_X(se_x, val_p, grupo_volatilidad, val_r)
se_y = serie_Y(se_x, val_r)
se_z = serie_Z(se_x, val_r, boton)

semillas = (se_z['VISCOSITY ABE'].iloc[-2], 
            se_z['VISCOSITY ABE'].iloc[-2] + 0.005, 
            se_z['VISCOSITY ABE'].iloc[-2] - 0.005)

# print('PREDICCIÓN 1:\n')
prediccion_1 = prediccion_velocidad(se_z, semillas[0])
# print('PREDICCIÓN 2:\n')
prediccion_2 = prediccion_velocidad(se_z, semillas[0], 0.005)
# print('PREDICCIÓN 3:\n')
prediccion_3 = prediccion_velocidad(se_z, semillas[0], -0.005)

# Volvemos a truncar la serie_X a la cantidad de datos
se_x = se_x.head(cantidad_datos)
se_z = se_z[['TIME',
            'RANGE',
            'DENSITY',
            'VELOCITY',
            'PRESSURE',
            'VISCOSITY',
            'VISCOSITY ABE',
            'SUBTRACTION VISCOSITIES']
            ].head(len(se_z)-1).round(decimals=8)


# CALCULOS
resumen = calculos_resumen_serie(se_x, se_z, val_p, val_r, semillas).round(decimals=6)
# print(resumen)



# =============================================================================
# ~                             TESTING
# =============================================================================

# # Notas:



# print('selfa:\n', se_x)
# print('se_y:\n', se_y)
# print('se_z:\n', se_z.to_string())



pinta([se_x], [['INITIAL SERIE']])
# pinta(se_y, ['normal'])
# pinta(se_z, ['val_prom'])

pinta([se_z], [['DENSITY']])
pinta([se_z], [['VELOCITY']])
pinta([se_z], [['PRESSURE']])

pinta([se_x], [['VOLATILITY']])
# # pinta(se_y, ['volatilidad'])
# # pinta(se_z, ['volatilidad'])

# # pinta(se_z, ['term_visc'])
pinta([se_z], [['VISCOSITY', 'VISCOSITY ABE']])

pinta([se_z], [['SUBTRACTION VISCOSITIES']])

pinta([se_z,
        prediccion_1, 
        prediccion_2, 
        prediccion_3], [['VELOCITY'], 
                        ['VELOCITY PREDICT'], 
                        ['VELOCITY PREDICT'], 
                        ['VELOCITY PREDICT']])


# Lista de tablas a exportar
# se_z['RANGE'] = [f"{i}-{i+25}" for i in range(0, cantidad_datos, val_r)]
lista_tablas=[se_x, se_y, se_z]

# RUTA DONDE CREAR EL EXCEL DE EXPORTACIÓN
# ruta_excel = r"G:\My Drive\A.Puente_SDG\CFD - Timoteo\Python"
ruta_doc = r"G:\.shortcut-targets-by-id\1QxyDLcOmaze0f2S9vwBTDDBIvubizoRL\A.Puente_SDG\CFD - Timoteo\Python"
nombre_doc = "DATAFRAME"+".docx"

directorio_exportacion = os.path.join(ruta_doc, nombre_doc)

# RUTA DONDE ESTA LA CARPETA AUXILIAR DE IMAGENES
# DEBERIA SER LA MISMA CARPETA EN LA QUE ESTÁ EL CÓDIGO
# MIRAR FUNCION PINTA, ES DONDE SE CREA LA CARPETA
directorio_script = os.getcwd()
carpeta_con_imagenes = os.path.join(directorio_script, 'temp_graficas')
# r'C:\Users\bogdan.stefan\Desktop\Temp_LOCAL\Carpeta temporal graficas'


# exportar_series(lista_tablas, directorio_exportacion, carpeta_con_imagenes)
export_ficha_word(resumen, se_z, carpeta_con_imagenes, hoja)







