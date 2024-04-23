import os
import docx
from docx.shared import Pt, Mm
import pandas as pd

def calculos_resumen_serie(serie_x, serie_z, val_p, val_r, seeds):
    tabla_resumen = pd.DataFrame()
    # print(serie_x)
    velocidad = serie_z['VELOCITY'].to_numpy()
    viscosidad = serie_z['VISCOSITY'].to_numpy()
    viscosidad_ABE = serie_z['VISCOSITY ABE'].to_numpy()
    substrac_visco  = serie_z['SUBTRACTION VISCOSITIES'].to_numpy()
    
    entropy = serie_x['ENTROPY'].to_numpy()
    ind_valp = serie_x['ind_valp'].to_numpy()
    # CALCULO IC
    max_aux = np.maximum(viscosidad_ABE, viscosidad)
    min_aux = np.minimum(viscosidad_ABE, viscosidad)
    IC = (np.nansum(abs(substrac_visco))/len(substrac_visco))/(
        np.nanmax(max_aux)- np.nanmin(min_aux))
    # print('IC: ', IC)
    
    # CALCULO SHANNON ENTROPY
    shannon_entropy = abs(np.sum(entropy))
    # print('SHANNON ENTROPY: ', shannon_entropy)
    
    # CALCULO SUBSTRACTION VISCOSITIES
    max_substract_visco = np.nanmax(substrac_visco)
    min_substract_visco = np.nanmin(substrac_visco)
    diff_max_min_visco = max_substract_visco - min_substract_visco
    # print(diff_max_min_visco, max_substract_visco, min_substract_visco)
    
    #CALULO TYPICAL DEVIATION VELOCITY
    std_velocity = np.nanstd(velocidad, ddof=1)
    # print('DESVIACION TIP VELOCIDAD: ', std_velocity)
    
    # CALCULO NODOS
    num_points = len(ind_valp)
    num_nodes = np.count_nonzero(ind_valp)
    # print(num_points, val_p, num_nodes)
    
    # CALCULOS AVERAGE SUBSTRAC VISCOSITIES
    mean_substrc_visco = np.nanmean(substrac_visco)
    
    # Calculo NUMERO DE INERVALOS
    num_intervalos = len(velocidad)
    
    tabla_resumen['IC'] = [IC]
    tabla_resumen['SHANNON ENTROPY'] = [shannon_entropy] 
    tabla_resumen['MAX-MIN (SUBSTRACTION VISCOSITIES)'] = [diff_max_min_visco]
    tabla_resumen['MAX (SUBSTRACTION VISCOSITIES)'] = [max_substract_visco]
    tabla_resumen['MIN (SUBSTRACTION VISCOSITIES)'] = [min_substract_visco]
    tabla_resumen['IS (STANDARD DEVIATION VELOCITY)'] = [std_velocity]
    tabla_resumen['AVERAGE SUBSTRACION VISCOSITIES'] = [mean_substrc_visco]
    tabla_resumen['SERIE NUMBER TOTAL'] = [num_points]
    tabla_resumen['\"p\" VALUE FOR NODES'] = [val_p]
    tabla_resumen['NODES NUMBER'] = [num_nodes]
    tabla_resumen['INTERVAL NUMBER'] = [num_intervalos]
    tabla_resumen['NUMBER POINTS EACH INTERVAL'] = [val_r]
    
    for i, seed in enumerate(seeds, 1):
        nom_columna = f'SEED - {i}'
        tabla_resumen[nom_columna] = seed
    
    print('RESUMEN\n\n', tabla_resumen.to_string())
    return tabla_resumen

            
def export_ficha_word(resumen_caracteristicas, serie_z, carpeta_con_imagenes, nombre_datos=''):
    
    doc = docx.Document()
    
    ajustar_margenes(doc, 10, 5, 0, 0)
    
    doc.add_heading(f'Ficha: {nombre_datos}', 0)
    
    # p = doc.add_paragraph('The ')
    # p.add_run('International System of Units').bold = True
    # p.add_run(', known by the international abbreviation ')
    # p.add_run('SI').bold = True
    # p.add_run(', is the modern form of the metric system ')
    # p.add_run('(Wikipedia: International System of Units)').italic = True
    
    # doc.add_heading('Properites', level=1)
    exportar_tabla_word(doc, resumen_caracteristicas)
    doc.add_paragraph('')
    
    parrafo = exportar_plots_word(doc, carpeta_con_imagenes)
    ajustar_interlineado_entre_parrgrafos(parrafo, 5)       
    # doc.add_heading('Table Serie Z', level=1)
    # exportar_tabla_word(doc, serie_z)
    
    # Settings
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Verdana'
    font.size = Pt(9)
    
    # Save the Word doc
    doc.save('Ficha_Serie.docx') 

def exportar_tabla_word(doc, tabla):
    # Initialise the table
    t = doc.add_table(rows=tabla.shape[1], cols=2)
    # t.style = 'List Table 2 Accent 1'
    t.autofit = True
    # print(type(t))
    # Add the column headings
    for j in range(tabla.shape[1]):
        titulo = tabla.columns[j].strip()
        valor = tabla.iloc[(0,j)]
        t.cell(j, 0).text = str(titulo)
        t.cell(j, 0).paragraphs[0].runs[0].bold = True
        t.cell(j, 1).text = str(valor)
    t.autofit = True
            
def exportar_tabla_word_bullet (doc, tabla):
        for columna in tabla.columns:
            p = doc.add_paragraph(style = 'List Bullet')
            # Añade el texto del título de la columna en negrita
            p.add_run(columna.strip() + ': ').bold = True
            # Añade el valor de la posición 0 sin negrita
            p.add_run(str(tabla[columna][0])).bold = False
            
        doc.add_paragraph()
        
def exportar_tabla_word_t(doc, tabla, transpuesto=False):
    # Initialise the table
    if transpuesto:
        t = doc.add_table(rows=tabla.shape[1]+1, cols=tabla.shape[0])
    else:
        t = doc.add_table(rows=1, cols=tabla.shape[1])

    # Add the column headings
    if transpuesto:
        for i in range(tabla.shape[1]):
            cell = tabla.index[i]
            t.cell(i+1, 0).text = str(cell)
            t.cell(i+1, 0).paragraphs[0].runs[0].bold = True
    else:
        for j in range(tabla.shape[1]):
            cell = tabla.columns[j].strip()
            t.cell(0, j).text = str(cell)
            t.cell(0, j).paragraphs[0].runs[0].bold = True

    # Add the body of the data frame
    for i in range(tabla.shape[0]):
        for j in range(tabla.shape[1]):
            if transpuesto:
                cell = tabla.iat[i, j]
                t.cell(j+1, i).text = str(cell)
            else:
                cell = tabla.iat[i, j]
                t.cell(i+1, j).text = str(cell)

            
def exportar_plots_word(doc, carpeta_con_imagenes):
    lista_imagenes = [archivo for archivo in os.listdir(carpeta_con_imagenes) if archivo.lower().endswith('.png')]
    paragraph = doc.add_paragraph()
    ajustar_interlineado_entre_parrgrafos(paragraph, 0.2)
    for i, imagen in enumerate(lista_imagenes, 1):
        ruta_imagen = os.path.join(carpeta_con_imagenes, imagen)
        if i % 3+ 1:
            paragraph.add_run().add_text('  ')
        run_imagen = paragraph.add_run()
        run_imagen.add_picture(ruta_imagen, width=Mm(63))  # Ajusta el ancho según sea necesario
        if i%3 == 0:
            paragraph = doc.add_paragraph()
            ajustar_interlineado_entre_parrgrafos(paragraph, 0.2)
    return paragraph


def ajustar_margenes(doc, margen_izquierdo=1, margen_derecho=1, margen_superior=1, margen_inferior=1):
    sections = doc.sections
    for section in sections:
        section.left_margin = Mm(margen_izquierdo)
        section.right_margin = Mm(margen_derecho)
        section.top_margin = Mm(margen_superior)
        section.bottom_margin = Mm(margen_inferior)
    
def cambiar_interlineado(parrafo, interlineado):
    for run in parrafo.runs:
        run.font.size = Pt(7)  # Ajusta el tamaño de fuente según sea necesario
        run._element.get_or_add_pPr().get_or_add_spacing().line = Pt(interlineado)
    
def ajustar_interlineado_entre_parrgrafos(parrafo, espacio_despues):
    parrafo.paragraph_format.space_after = Pt(espacio_despues)